"""
Office, Spreadsheet, Tabular, and Structured Document Loaders
Handles .docx, .doc, .xlsx, .xls, .csv, .tsv, .json, .jsonl, .html, .htm, .xml, .yaml, .yml
"""

import csv
import json
import re
from pathlib import Path

from ingestion.base_loader import BaseLoader
from utils.helpers import normalize_text
from utils.logger import logger
from utils.models import Document


class DocxLoader(BaseLoader):
    """Load Microsoft Word .docx and .doc files."""

    def load(self, source: str) -> list[Document]:
        path = Path(source).resolve()
        self._validate_file(path)
        logger.info(f"Loading Word document: {path.name}")

        text_parts = []

        try:
            import docx

            doc = docx.Document(str(path))
            for p in doc.paragraphs:
                if p.text.strip():
                    text_parts.append(p.text.strip())

            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(
                        cell.text.strip() for cell in row.cells if cell.text.strip()
                    )
                    if row_text:
                        text_parts.append(row_text)
        except Exception as e:
            logger.warning(
                f"python-docx failed for {path.name}, attempting text fallback: {e}"
            )
            try:
                raw = path.read_text(encoding="utf-8", errors="ignore")
                printable = "".join(c for c in raw if c.isprintable() or c in "\n\r\t")
                text_parts = [printable]
            except Exception as fallback_err:
                logger.error(f"Docx loading failed for {path.name}: {fallback_err}")
                raise RuntimeError(f"Word loading failed for {path.name}: {e}") from e

        text = normalize_text("\n\n".join(text_parts))

        if not text:
            logger.warning(f"Word document {path.name} is empty")
            return []

        return [
            Document(
                content=text,
                source=path.name,
                metadata={
                    "source": path.name,
                    "full_path": str(path),
                    "page": 1,
                    "file_type": path.suffix.lstrip(".").lower(),
                },
            )
        ]


class ExcelLoader(BaseLoader):
    """Load Excel .xlsx and .xls files."""

    def load(self, source: str) -> list[Document]:
        path = Path(source).resolve()
        self._validate_file(path)
        logger.info(f"Loading Excel file: {path.name}")

        sheet_texts = []

        try:
            import pandas as pd

            excel_file = pd.ExcelFile(str(path))
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(excel_file, sheet_name=sheet_name)
                df = df.dropna(how="all")
                if not df.empty:
                    sheet_str = f"--- Sheet: {sheet_name} ---\n" + df.to_string(
                        index=False
                    )
                    sheet_texts.append(sheet_str)
        except Exception as e:
            logger.warning(
                f"Pandas excel load failed for {path.name}, trying openpyxl: {e}"
            )
            try:
                import openpyxl

                wb = openpyxl.load_workbook(str(path), data_only=True)
                for sheet_name in wb.sheetnames:
                    sheet = wb[sheet_name]
                    rows = []
                    for row in sheet.iter_rows(values_only=True):
                        row_vals = [str(v) for v in row if v is not None]
                        if row_vals:
                            rows.append(" | ".join(row_vals))
                    if rows:
                        sheet_texts.append(
                            f"--- Sheet: {sheet_name} ---\n" + "\n".join(rows)
                        )
            except Exception as fallback_err:
                logger.error(f"Excel loading failed for {path.name}: {fallback_err}")
                raise RuntimeError(f"Excel loading failed for {path.name}: {e}") from e

        text = normalize_text("\n\n".join(sheet_texts))

        if not text:
            logger.warning(f"Excel file {path.name} is empty")
            return []

        return [
            Document(
                content=text,
                source=path.name,
                metadata={
                    "source": path.name,
                    "full_path": str(path),
                    "page": 1,
                    "file_type": path.suffix.lstrip(".").lower(),
                },
            )
        ]


class CsvLoader(BaseLoader):
    """Load CSV and TSV files into searchable tabular text format."""

    def load(self, source: str) -> list[Document]:
        path = Path(source).resolve()
        self._validate_file(path)
        logger.info(f"Loading CSV/TSV file: {path.name}")

        delimiter = "\t" if path.suffix.lower() == ".tsv" else ","

        lines = []
        try:
            with open(path, encoding="utf-8", errors="replace") as f:
                reader = csv.reader(f, delimiter=delimiter)
                headers = None
                for i, row in enumerate(reader):
                    if not row:
                        continue
                    if i == 0:
                        headers = row
                        lines.append("Columns: " + " | ".join(headers))
                    else:
                        if headers and len(row) == len(headers):
                            pair_str = ", ".join(
                                f"{h}: {v}" for h, v in zip(headers, row) if v.strip()
                            )
                            lines.append(f"Row {i}: {pair_str}")
                        else:
                            lines.append(" | ".join(row))
        except Exception as e:
            logger.error(f"CSV loading failed for {path.name}: {e}")
            raise RuntimeError(f"CSV loading failed for {path.name}: {e}") from e

        text = normalize_text("\n".join(lines))

        if not text:
            logger.warning(f"CSV file {path.name} is empty")
            return []

        return [
            Document(
                content=text,
                source=path.name,
                metadata={
                    "source": path.name,
                    "full_path": str(path),
                    "page": 1,
                    "file_type": path.suffix.lstrip(".").lower(),
                },
            )
        ]


class JsonLoader(BaseLoader):
    """Load .json and .jsonl files."""

    def load(self, source: str) -> list[Document]:
        path = Path(source).resolve()
        self._validate_file(path)
        logger.info(f"Loading JSON file: {path.name}")

        ext = path.suffix.lower()
        content_lines = []

        try:
            raw = path.read_text(encoding="utf-8", errors="replace")
            if ext == ".jsonl":
                for line in raw.splitlines():
                    line = line.strip()
                    if line:
                        obj = json.loads(line)
                        content_lines.append(json.dumps(obj, indent=2))
            else:
                data = json.loads(raw)
                content_lines.append(json.dumps(data, indent=2))
        except Exception as e:
            logger.error(f"JSON loading failed for {path.name}: {e}")
            raise RuntimeError(f"JSON loading failed for {path.name}: {e}") from e

        text = normalize_text("\n\n".join(content_lines))

        if not text:
            logger.warning(f"JSON file {path.name} is empty")
            return []

        return [
            Document(
                content=text,
                source=path.name,
                metadata={
                    "source": path.name,
                    "full_path": str(path),
                    "page": 1,
                    "file_type": "json",
                },
            )
        ]


class HtmlLoader(BaseLoader):
    """Load .html and .htm files into clean text."""

    def load(self, source: str) -> list[Document]:
        path = Path(source).resolve()
        self._validate_file(path)
        logger.info(f"Loading HTML file: {path.name}")

        try:
            raw = path.read_text(encoding="utf-8", errors="replace")
            # Strip script and style tags
            raw = re.sub(
                r"<(script|style).*?>.*?</\1>", "", raw, flags=re.DOTALL | re.IGNORECASE
            )
            # Strip remaining HTML tags
            text = re.sub(r"<[^>]+>", " ", raw)
            text = normalize_text(text)
        except Exception as e:
            logger.error(f"HTML loading failed for {path.name}: {e}")
            raise RuntimeError(f"HTML loading failed for {path.name}: {e}") from e

        if not text:
            logger.warning(f"HTML file {path.name} is empty")
            return []

        return [
            Document(
                content=text,
                source=path.name,
                metadata={
                    "source": path.name,
                    "full_path": str(path),
                    "page": 1,
                    "file_type": "html",
                },
            )
        ]
